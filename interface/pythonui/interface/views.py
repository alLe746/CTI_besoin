from datetime import datetime

import docx
import json
import pandas as pd
from django.shortcuts import render, redirect
from django.http import HttpResponse, JsonResponse
# Create your views here.
from .DocumentMeta import DocumentMeta,FileUpload,Modeldocx,Jsonchoice,FileUploadjson
from docx import Document
import os

def index(request):

    return render(request, "interface/index.html")



def wordtojson(request):
    form = FileUpload(request.POST)
    if request.method == 'POST':
        document = request.FILES['file']
        with open('tempword.docx', 'wb+') as destination:
            for chunk in document.chunks():
                destination.write(chunk)
        docu_valid = True
        try:
            document = Document('tempword.docx')
        except docx.opc.exceptions.PackageNotFoundError as e:
            print("document introuvable, vérifier l'emplacement")
            print("message d'erreur : " + str(e))
            docu_valid = False
        tables = []
        for table in document.tables:
            df = [['' for i in range(len(table.columns))] for j in range(len(table.rows))]
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    if cell.text:
                        df[i][j] = cell.text
            tables.append(pd.DataFrame(df))

        res = {}
        # Partie demandeur
        for i in range(1, 5):
            if i == 3:
                res[tables[1][0][i]] = datetime.strptime(tables[1][1][i], "%d/%m/%Y").strftime("%Y-%m-%d")
                res['@timestamp'] = res[tables[1][0][i]]
            else:
                res[tables[1][0][i]] = tables[1][1][i]
        for i in range(1, 7):
            if i == 5:
                if tables[2][1][i] != "":
                    res[tables[2][0][i]] = datetime.strptime(tables[2][1][i], "%d/%m/%Y").strftime("%Y-%m-%d")
            else:
                res[tables[2][0][i]] = tables[2][1][i]

        for i in range(1, 3):
            res[tables[3][0][i]] = tables[3][1][i]
        # Partie Reponse

        for i in range(1, 6):
            if i==3 or i==4:
                if tables[4][1][i] != "":
                    res[tables[4][0][i]] = datetime.strptime(tables[4][1][i], "%d/%m/%Y").strftime("%Y-%m-%d")
            else:
                res[tables[4][0][i]] = tables[4][1][i]

        response = HttpResponse(json.dumps(res),
            content_type='application/json')
        response['Content-Disposition'] = 'attachment; filename=download.json'
        return response
    else:
        return render(request, "interface/converttojson.html", {'form':form})

def jsontoword(request):
    form = FileUploadjson(request.POST)
    if request.method == 'POST':
        document = request.FILES['file']
        json_return = ""

        with open('tempjson.json', 'wb+') as destination:
            for chunk in document.chunks():
                destination.write(chunk)

        docu_valid = True
        file=open("tempjson.json")
        res=json.load(file)
        file.close()
        num_table = 0
        document = Document(os.path.join(os.path.dirname(__file__), 'Template demande de besoin.docx'))
        for table in document.tables:
            if num_table != 0:
                num_row = 0
                for i, row in enumerate(table.rows):
                    if num_row != 0:
                        temp = ""
                        for j, cell in enumerate(row.cells):
                            if (cell.text == "") and (temp != ""):
                                if temp=="Date de la demande" or temp=="Date de livraison" or temp=="Date de réception de la demande" or temp=="Date de traitement":
                                    cell.text = datetime.strptime((res[temp]),"%Y-%m-%d").strftime("%d/%m/%Y")
                                else:
                                    cell.text = res[temp]
                                # print("json: " + json_retour[temp])
                            if cell.text:
                                temp = cell.text
                                # print("cell.text: "+cell.text)
                    num_row = num_row + 1
            num_table = num_table + 1
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=download.docx'
        document.save(response)
        return response
    else:
        return render(request, "interface/converttoword.html", {'form': form})

def ressources(request):
    if request.method=='POST':
        document = Document(os.path.join(os.path.dirname(__file__), 'Template demande de besoin.docx'))
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=download.docx'
        document.save(response)
        return response
    else:
        return render(request, "interface/ressources.html")